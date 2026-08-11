from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).with_name("OPTICSLAB_PHYSICS_AND_SCHEMATIC_REDESIGN_PLAN.docx")
BLUE = "2E74B5"
NAVY = "1F4D78"
PALE_BLUE = "E8EEF5"
GRAY = "666666"
PAGE_WIDTH_DXA = 9360


def set_run_font(run, size=11, bold=None, color="000000"):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table_pr = table._tbl.tblPr
    tbl_w = table_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        table_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    table_pr.append(tbl_ind)
    grid = table._tbl.tblGrid
    for index, width in enumerate(widths):
        grid.gridCol_lst[index].set(qn("w:w"), str(width))
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_mar = OxmlElement("w:tcMar")
            for side in ("top", "start", "bottom", "end"):
                node = OxmlElement(f"w:{side}")
                node.set(qn("w:w"), "120" if side in ("start", "end") else "80")
                node.set(qn("w:type"), "dxa")
                tc_mar.append(node)
            cell._tc.get_or_add_tcPr().append(tc_mar)


def write_cell(cell, text, header=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(text)
    set_run_font(run, size=9.5, bold=header, color="000000")
    if header:
        set_cell_shading(cell, PALE_BLUE)


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    paragraph.paragraph_format.space_after = Pt(7 if level == 1 else 5)
    run = paragraph.add_run(text)
    set_run_font(run, size=16 if level == 1 else 13, bold=True, color=BLUE if level == 1 else NAVY)
    return paragraph


def add_body(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.25
    set_run_font(paragraph.add_run(text), size=11)
    return paragraph


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.25
        set_run_font(paragraph.add_run(item), size=10.5)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for index, value in enumerate(headers):
        write_cell(table.rows[0].cells[index], value, header=True)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            write_cell(cells[index], value)
    return table


def add_page_footer(section):
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(footer.add_run("OpticsLab | Physics and Schematic Redesign Plan | 2026-08-10"), size=8, color=GRAY)


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    add_page_footer(section)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    set_run_font(title.add_run("OpticsLab Physics and Schematic Redesign Plan"), size=22, bold=True, color="000000")
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    set_run_font(subtitle.add_run("Phase 0 baseline, architecture contract, verification matrix, and delivery gates"), size=12, color=GRAY)
    add_table(doc, ["Document", "Value"], [
        ("Scope", "Physics interaction, component semantics, BeamGraph, bench-to-schematic projection, and export QA"),
        ("Product boundary", "2D geometric/paraxial workbench and research-schematic draft tool; not a commercial optical-design replacement"),
        ("Status", "Living implementation contract; Phase 0 established before production-code changes"),
        ("Source basis", "README, rescue plan, upgrade roadmap, adapter and validation documentation, Phase 2 design/plan, source, and tests")
    ], [2500, 6860])

    add_heading(doc, "1. Product Objective and Boundary")
    add_body(doc, "The target product is a trustworthy 2D optics workbench: users can inspect geometric paths, create ideal thin-lens teaching diagrams, and turn a bench scene into an editable schematic draft. Every result must expose its model boundary before a user can mistake it for a quantitative design result.")
    add_bullets(doc, [
        "exact_geometric: a deterministic geometric law with golden coverage.",
        "paraxial_approximation: ideal thin-lens or small-angle behavior with visible assumptions.",
        "educational_visualization or experimental: useful only within stated limits.",
        "visual_only: semantic symbol without a physics claim; UI and exports must preserve that label."
    ])

    add_heading(doc, "2. Current Reproduction Record")
    add_table(doc, ["Issue", "Observed risk", "Reproduction contract", "Exit condition"], [
        ("A: lens rotation", "ThinLens derives visual curvature and label from focalLength sign while geometry uses a view-oriented axis. Thick-lens front/back radii are not represented as immutable local-surface semantics.", "Create failing 0/90/180/270 degree tests for biconvex, biconcave, plano-convex, and plano-concave models; trace equivalent rays in local coordinates.", "A 180 degree rotation changes pose and incidence only. Biconvex never becomes biconcave, and focal power remains invariant."),
        ("B: schematic conversion", "Projection is currently a separate adapter path. It must preserve independent placement records and never infer optical paths from rendered stroke color.", "Create a failing bench-to-schematic conversion test with placements, rotation, manual schematic override, and BeamGraph branches.", "Initial projection uses one documented transform; later projections fill only missing placements and retain manual schematic layout.")
    ], [1400, 2700, 2600, 2660])

    add_heading(doc, "3. Physics Coordinates and Pose Invariants")
    add_body(doc, "All component interaction must be evaluated in component-local space and converted back to world space. Angle changes are rigid transforms; they never mutate type, material, focal length, curvature, or port identity.")
    add_table(doc, ["Concept", "Contract"], [
        ("World frame", "Canvas coordinates used for scene placement, rendering, and persisted bench placement."),
        ("Local frame", "Origin at component center; local +u follows optical axis; local +v follows component aperture/height direction."),
        ("Ports", "Stable IDs are local-space anchors with local directions. World position and normal are derived using the pose transform."),
        ("Surface normal", "Derived from local surface geometry at the hit point, then transformed to world space; never inferred from an SVG path direction."),
        ("Lens identity", "type, shapeProfile, material, focalLength, curvature, and port IDs are immutable under rotation. 180 degrees changes front-facing side only."),
        ("Transform API", "rayToLocalSpace, rayToWorldSpace, intersectComponent, resolveSurfaceNormal, and interactWithComponent are pure functions or thin adapters around pure functions.")
    ], [2200, 7160])

    add_heading(doc, "4. Unified Optical Interaction Interface")
    add_body(doc, "The rendering/UI layer asks the interaction layer for structured outcomes. It does not classify a component from a canvas fill, visual arrow, or SVG stroke.")
    add_table(doc, ["Function", "Required output"], [
        ("rayToLocalSpace / rayToWorldSpace", "Ray origin and normalized direction under an explicit pose transform."),
        ("intersectComponent / resolveSurfaceNormal", "Hit point, local and world normal, surface ID, input port candidate, and finite distance."),
        ("interactWithComponent", "One or more outcome rays with output port, interaction type, energy rule, and termination state."),
        ("propagatePolarization / propagateWavelength / propagateFrequency", "Explicit state transition with reliability label and model assumption."),
        ("splitBeam / terminateBeam", "Structured child/terminal records; no visual convention is used as data."),
        ("buildInteractionMetadata", "source ray ID, parent segment ID, component ID, ports, normal, wavelength, frequency, shift, polarization, intensity, path kind, and termination reason.")
    ], [3000, 6360])

    add_heading(doc, "5. Component Reliability and Test Matrix")
    add_table(doc, ["Family", "Representative components", "Baseline semantics", "Minimum golden coverage"], [
        ("Sources", "Laser, point, line, fan, LED", "Emission geometry; wavelength/polarization only where modeled", "Port direction, normalization, metadata source ID"),
        ("Lens", "Thin, cylindrical, aspheric, GRIN", "Thin lens is paraxial; others remain experimental until proven", "0/90/180/270 pose, forward/reverse, local/world round-trip, power invariant"),
        ("Reflective/splitting", "Mirror, curved mirror, splitter, PBS, dichroic", "Reflection and branch output ports", "Normal law, branch intensity, reverse path, PBS polarization branch"),
        ("Polarization/modulation", "Wave plates, polarizer, Faraday, AOM/EOM", "Jones-state or documented approximation", "Wave plate transform, polarization selection, AOM frequency shift"),
        ("Passive/detection", "Aperture, fiber, screen, photodiode, atomic cell", "Termination, attenuation, or visual-only semantics", "Termination reason, detector accumulation, UI reliability display"),
        ("Diagram-only", "Imported or unsupported components", "visual_only", "No precise-physics wording; symbol and export metadata preserved")
    ], [1500, 2200, 2700, 2960])

    add_heading(doc, "6. OpticsDocument v3 and BeamGraph")
    add_body(doc, "OpticsDocument v3 owns canonical component identity, BeamGraph connectivity, and independent bench/schematic view state. Bench and schematic positions never overwrite one another. BeamGraph is the only source for optical topology in schematic paths.")
    add_bullets(doc, [
        "BeamGraph nodes carry component ID/type/name, stable port descriptors, and serializable properties.",
        "Edges carry component/port endpoints and structured semantics for reflection, split, return, frequency shift, auxiliary/dashed path, and termination.",
        "A dashed path is represented by pathKind or auxiliary metadata, never inferred from color.",
        "Round-trip serialization preserves schema version, canonical components, graph edges, both placements, manual-layout lock state, annotations, and reliability metadata."
    ])

    add_heading(doc, "7. Bench-to-Schematic Preservation Policy")
    add_table(doc, ["Rule", "Implementation constraint"], [
        ("Initial projection", "Apply exactly one recorded coordinate transform from bench placement to schematic placement; retain center, relative distance, angle, and connectivity."),
        ("Manual layout", "Existing schematic placement is authoritative. SchematicProjector may create only missing records and must not overwrite moved symbols."),
        ("Cross-view updates", "Add/remove/rename/connect synchronize through canonical components and BeamGraph. Placement values remain view-local."),
        ("Path source", "Generate schematic paths from BeamGraph and interaction metadata, including reflection, branches, return paths, frequency shift, and auxiliary paths."),
        ("Idempotence", "Re-running conversion on the same document changes no manually edited schematic placement or symbol ID.")
    ], [2500, 6860])

    add_heading(doc, "8. Semantic SVG Symbol System")
    add_body(doc, "Symbols use native SVG DOM groups. No raster dependency, screenshot trace, or universal rectangle placeholder is permitted for supported optical classes.")
    add_bullets(doc, [
        "Each symbol has a stable group id plus data-component-id and data-component-type.",
        "Required layers separate symbol body, ports, label anchor, interaction markers, and selection affordances.",
        "Initial symbol set: laser, thin/convex/concave/plano lens variants, mirror, beam splitter, PBS, wave plate, polarizer, AOM, prism, grating, fiber, atomic cell, screen, detector, and terminator.",
        "Symbols remain selectable, editable, movable, and serializable after SVG export/import.",
        "Use original maintainable SVG symbols. No external asset is admitted without verified repository URL, license, redistribution terms, version, and modification record."
    ])

    add_heading(doc, "9. Phased Implementation and Gates")
    add_table(doc, ["Phase", "Deliverable", "Gate"], [
        ("0", "This DOCX, source audit, actual reproductions, baseline screenshots/logs", "No production refactor before both issue records are executable."),
        ("1", "Pose model and pure interaction boundary", "Failing coordinate/rotation tests turn green; local/world round-trip is covered."),
        ("2", "Lens and component interaction fixes", "Golden tests cover lens rotation, mirror/PBS branches, polarization, frequency, and termination."),
        ("3", "Stable ports and BeamGraph metadata", "Branch, bidirectional, frequency-shift, auxiliary, and terminal graph semantics persist."),
        ("4", "Position-preserving projector and semantic SVG symbols", "Symbol snapshots and conversion tests prove no manual layout overwrite."),
        ("5", "Dual-view sync, save, restore, re-import", "Round-trip tests and browser flows retain independent placements."),
        ("6", "Export, visual QA, 100-component performance, legacy audit", "Desktop/mobile screenshots, console clean, SVG/PNG/PDF structure stable."),
    ], [800, 4700, 3860])

    add_heading(doc, "10. Test, Screenshot, Performance, and Release Gates")
    add_bullets(doc, [
        "Every behavior change follows RED -> expected failure -> minimal GREEN -> focused tests -> full unit suite -> relevant Playwright workflow -> desktop and mobile screenshot inspection -> console review.",
        "Required suites include LensRotationInvariant, ComponentInteractionGolden, ComponentPorts, BeamGraph, SchematicProjector, SchematicSymbols, SchematicPositionPreservation, DualViewSync, OpticsDocumentRoundTrip, and four named e2e workflows.",
        "Coverage scenario: Laser -> AOM -> Mirror -> PBS -> WavePlate -> AtomicCell, including reflection/return, PBS branches, frequency shift, polarization, visual-only or terminal atomic-cell semantics, and position preservation.",
        "Release cannot claim professionalization until unit and e2e suites pass, screenshots show no blank/overlap/mis-semantic symbol, exports are structurally stable, and all physical limits are visible in UI and documentation."
    ])

    add_heading(doc, "11. Phase 0 Evidence and Remaining Risks")
    add_table(doc, ["Evidence", "Result"], [
        ("Baseline unit suite", "npm test passed before the new reproduction contracts were added."),
        ("Baseline browser smoke", "npm run test:e2e passed. The canvas screenshot is stored at output/playwright/smoke-canvas.png and the smoke workflow reported no captured console errors."),
        ("Issue A reproduction", "tests/unit/LensRotationInvariant.test.js failed because ThinLens has no shapeProfile. The existing model cannot encode immutable biconvex/biconcave surface identity across a 180 degree pose change."),
        ("Issue B reproduction", "tests/unit/SchematicPositionPreservation.test.js failed because SchematicProjector replaced bench placement {160,280,15} with automatic placement {180,120,0}."),
        ("Phase 1 implementation", "src/physics/ComponentInteraction.js now provides pure local/world transforms, normal resolution, component interaction delegation, state propagation, split/termination records, and traceable metadata. ThinLens serializes shapeProfile; first projection copies valid bench placement only for missing schematic entries."),
        ("Phase 1 verification", "Focused tests, npm test, and npm run test:e2e passed after implementation. The updated smoke canvas was inspected and console capture reported no errors."),
        ("Phase 2 lens golden test", "LensRotationInvariant covers biconvex central rays at 0/90/180/270 degrees in both propagation directions. It verifies incident-facing normals, preserved transmission direction, focal length, and immutable surface identity under the current paraxial model."),
        ("Phase 2 AOM golden test", "AOMInteractionGolden verifies direct AOM interaction emits zeroOrder and firstOrder branches with cumulative RF frequency metadata before RayTracer or BeamGraph post-processing. Energy split and geometric diffraction remain unchanged."),
        ("Phase 2 AtomicCell semantics", "AtomicCell is explicitly visual_only. AtomicCellSemantics verifies that it terminates a ray with visual_only_atomic_cell and emits no pseudo-physical Beer-Lambert transmission until a separately validated atom-light model exists."),
        ("Phase 2 visual-only detectors", "VisualOnlyDetectorSemantics verifies that CCDCamera and Spectrometer retain their UI aggregate data, emit no downstream ray, and terminate with explicit visual-only reasons rather than physical absorption claims."),
        ("Phase 2 visual-only annotations", "VisualOnlyAnnotationSemantics verifies that MagneticCoil and CustomComponent never intersect, block, or mutate optical rays. Their scope remains annotation only."),
        ("Phase 3 BeamGraph metadata", "Ray trace records and BeamGraph edges now carry interactionType, surfaceId, and an incident-facing unit surfaceNormal. BeamGraph verifies this both from a direct trace record and a live Laser-to-AOM RayTracer result; the AOM input normal no longer reverses after its incoming-face check."),
        ("Phase 4 semantic schematic", "Schematic projection refreshes BeamGraph branch, auxiliary, frequency, interaction, surface, and normal data while retaining manual path geometry and style. SVG paths expose those fields through data attributes. The native SVG symbol library now gives convex/concave lenses, splitter/PBS, polarizer, prism, grating, fiber, screen, AOM, mirror, wave plate, and atomic cell stable semantic forms without third-party assets."),
        ("Phase 5 round trip", "OpticsDocumentSerializer round-trips BeamGraph and schematic path interaction metadata together, including an AOM first-order frequency shift, auxiliary style, surface normal, and manually edited path points. Existing document file-controller and recovery tests remain green."),
        ("Phase 6 termination projection", "BeamGraph termination nodes now project as derived schematic terminationNodes with a dedicated SVG endpoint and reason label. Termination edges remain visible while the user component collection stays unchanged; these display nodes are rebuilt from BeamGraph rather than treated as editable optical components."),
        ("Phase 7 PowerMeter semantics", "PowerMeterSemantics verifies that the educational display meter accepts rays from either side with incident-facing normals, accumulates total/peak/average readings, terminates the incident ray, and resets its display state. It remains a teaching visualization rather than a calibrated instrument model."),
        ("Phase 8 detector and atomic semantics", "PolarizationAnalyzerSemantics verifies incident-facing normals from both directions, intensity-weighted linear Stokes accumulation, reset behavior, the analyzer termination reason, and agreement between the project right-circular Jones convention and its positive S3 / circular-right display label. AtomicCell is consistently visual-only: it terminates an incoming ray and its BeamGraph contract exposes only an input port, so diagrams cannot imply unsupported atomic-cell transmission."),
        ("Phase 9 thick-lens surfaces", "ThickLensSurfacePhysics replaces the previous single-plane thick-lens shortcut with finite-aperture spherical or planar front/back surface intersections. Each crossing uses Snell refraction with an explicit medium-index transition, preserves one quality factor at the physical exit, exposes front_surface/back_surface metadata, and keeps the internal lens edge out of the round-trip path category."),
        ("Phase 10 Fresnel branches", "Each thick-lens surface now emits scalar unpolarized Fresnel transmitted and reflected branches. The primary transmitted path carries the two real interface transmission factors and one lens-quality loss; internally reflected rays retain the lens medium. ThinLens ports distinguish transmitted and reflected graph branches while preserving the legacy output port for thin-lens paths."),
        ("Phase 11 SVG state semantics", "Schematic projection and SVG export now preserve BeamGraph wavelength, polarization, and intensity together with branch kind, frequency offset, interaction type, and surface metadata. These fields are emitted as structured data attributes, so a schematic remains inspectable without changing manually edited route geometry or path style."),
        ("Phase 12 thick-lens persistence", "ThinLens now owns a standard fromJSON restoration path. Custom thick-lens type, thickness, front/back radii, material dispersion, quality, and compatible Infinity/null values survive both component JSON and generic scene serialization. The legacy main-application loading branches delegate to this same restoration boundary instead of rebuilding a thin-lens default."),
        ("Phase 13 variable attenuation", "VariableAttenuatorSemantics verifies incident-facing normals from both directions, scalar intensity transmission, wavelength and Jones-state preservation, explicit attenuation termination, optical-density/dB readouts, and the configured transmission range clamp. Its reliability remains educational because the model intentionally excludes polarization-dependent and spectral attenuation."),
        ("Phase 14 point-source emission", "OpticalSources verifies PointSource centered uniform angular sampling, origin and wavelength preservation, total-intensity conservation, disabled-state behavior, and normalization of non-integer ray-count inputs. Its educational model is explicitly two-dimensional and does not claim three-dimensional radiometric behavior."),
        ("Phase 15 LED source and persistence", "OpticalSources verifies LED Gaussian wavelength sampling, endpoint fan geometry, random phase, unpolarized output, intensity conservation, disabled behavior, integer ray-count normalization, and generic scene round-trip for both PointSource and LEDSource. PointSource and LEDSource now own fromJSON restoration boundaries; legacy loading accepts numRays with rayCount compatibility. The model remains educational and does not claim package optics or calibrated radiometry."),
        ("Phase 16 pulsed laser source", "OpticalSources verifies PulsedLaserSource pulse energy, average power, transform-limited bandwidth, peak-power ray splitting, disabled behavior, integer ray-count normalization, and generic scene persistence. PulsedLaserSource now owns fromJSON restoration and legacy loading uses it. Reliability remains experimental because the runtime does not implement time-resolved propagation or chirp physics."),
        ("Phase 17 ring mirror", "RingMirror verifies center-hole transmission, outer-ring reflection with scalar 0.99 loss and phase reversal, rotated incident-facing normals, annular region boundaries, and generic scene persistence. RingMirror now owns fromJSON restoration. Its reliability remains educational because coating spectra and diffraction are not modeled."),
        ("Phase 18 dichroic mirror", "DichroicMirror verifies bounded sigmoid wavelength reflectivity, short-wave reflection, long-wave transmission, incident-facing normals, non-increasing branch energy, and generic scene persistence. DichroicMirror now owns fromJSON restoration. Its reliability remains experimental because coating-stack angle/polarization spectra are not modeled."),
        ("Phase 19 metallic mirror", "MetallicMirror verifies bounded reflectivity and finite phase shift for every supported metal across representative incidence angles, incident-facing normals, reflected direction, termination, and generic scene persistence. MetallicMirror now owns fromJSON restoration. Reliability remains experimental because its metal constants are single-wavelength values with simplified angular behavior."),
        ("Phase 20 electro-optic modulator", "ElectroOpticModulator verifies scalar phase modulation, cosine-square amplitude transmission, quality loss, axial transmission, and generic scene persistence of voltage/frequency/state. It now owns fromJSON restoration. Reliability remains experimental because crystal tensors, polarization coupling, and time-domain RF behavior are not modeled."),
        ("Phase 21 optical chopper", "OpticalChopper verifies blade-sector blocking, phase-rotated slot transmission, period, average duty-cycle transmission, normalized phase, and generic scene persistence. OpticalChopper now owns fromJSON restoration. Reliability remains educational because blade diffraction and phase-locked mechanics are not modeled."),
        ("Phase 22 Fabry-Perot cavity", "FabryPerotCavity verifies bounded Airy transmission, resonance peaks, finesse, free spectral range, linewidth, nearby resonance wavelengths, axial transmission, and generic scene persistence. It now owns fromJSON restoration. Reliability remains experimental because mode structure, mirror phase, and polarization-resolved resonance are not modeled."),
        ("Phase 23 white-light source", "OpticalSources verifies WhiteLightSource fast-mode sampling, full weighted-spectrum wavelength coverage, total-intensity conservation, integer ray-count normalization, and generic scene persistence. WhiteLightSource now owns fromJSON restoration. Its reliability remains educational because the spectrum is a fixed weighted visible approximation rather than calibrated spectral power data."),
        ("Phase 24 advanced-lens persistence", "AdvancedLenses verifies AsphericLens and GRINLens generic scene persistence together with their existing central/off-axis paraxial, GRIN profile, pitch, and symmetry contracts. Both now own fromJSON restoration. They remain partial/experimental because these tests do not validate full surface aberration or broad reference-case GRIN propagation."),
        ("Phase 25 unified legacy restoration", "The main legacy/preset loader now delegates WhiteLightSource, DichroicMirror, MetallicMirror, RingMirror, AsphericLens, GRINLens, ElectroOpticModulator, OpticalChopper, and FabryPerotCavity to their standard fromJSON boundaries. This removes stale constructor-order drift and preserves their serialized physical parameters consistently with generic scene deserialization."),
        ("Phase 26 placement and product QA", "The UI placement preview and final placement paths now construct AsphericLens with a zero conic constant and GRINLens with n0=1.6 and g=0.01, both at 90 degrees. Browser smoke coverage verifies their runtime defaults alongside a laser, mirror, and thick biconvex lens. The smoke runner now splits long Playwright snippets into short same-session calls to stay below the Windows command-line limit."),
        ("Phase 27 advanced-lens reference validation", "AdvancedLenses now verifies spherical and parabolic aspheric sag/slope references. GRINLens now selects the first physical face from the ray travel direction and propagates with a direction-aware local axis; bidirectional quarter-pitch transfer verifies the entry face, exit position, and analytic output slope. Both components remain partial/experimental because this is not full-aperture aberration analysis, full inhomogeneous-medium tracing, or calibrated lens validation."),
        ("DOCX QA", "Structural inspection passed: title, 36 paragraphs, 9 tables, and one section. Rendered PNG QA is blocked because no LibreOffice/soffice executable is installed in the environment."),
        ("Residual risk", "Thick lenses now trace two scalar geometric surfaces with unpolarized Fresnel branches, but they do not model thin-film coating spectra, polarization-resolved Fresnel coefficients, or higher-order aspheric aberration corrections. AtomicCell remains intentionally visual-only until a separately validated absorption/dispersion model exists.")
    ], [2500, 6860])

    add_heading(doc, "12. Known Limits and Residual Risks")
    add_table(doc, ["Risk", "Treatment"], [
        ("Advanced lens models", "Aspheric, GRIN, thick-surface, aberration, and coating behavior remain experimental until local-surface tests and validations exist."),
        ("Unsupported components", "Must carry visual_only or experimental metadata. Their UI wording cannot imply a precise model."),
        ("Schematic fidelity", "Schematic is an editable semantic representation, not a second physics engine. It consumes canonical data and BeamGraph."),
        ("External icon assets", "No imported asset is used until its license and redistribution conditions are recorded in source documentation."),
        ("Performance", "The 100-component target requires measured browser evidence after the new projector and symbol layer are integrated.")
    ], [2500, 6860])

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_document()
    print(OUTPUT)
